#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Extracts annotations from a PDF file to generate a response sheet for revision.
"""

import sys, io, argparse, subprocess
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.layout import LAParams, LTContainer, LTAnno, LTChar, LTTextBox
from pdfminer.converter import TextConverter
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
import pdfminer.pdftypes as pdftypes
import pdfminer.settings
import pdfminer.utils
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import csv

pdfminer.settings.STRICT = False

SUBSTITUTIONS = {
    u'ﬀ': 'ff',
    u'ﬁ': 'fi',
    u'ﬂ': 'fl',
    u'ﬃ': 'ffi',
    u'ﬄ': 'ffl',
    u'‘': "'",
    u'’': "'",
    u'“': '"',
    u'”': '"',
    u'…': '...',
}

ANNOT_SUBTYPES = frozenset({'Text', 'Highlight', 'Squiggly', 'StrikeOut', 'Underline'})

COLUMNS_PER_PAGE = 2 # default only, changed via a command-line parameter

DEBUG_BOXHIT = False

def boxhit(item, box):
    (x0, y0, x1, y1) = box
    assert item.x0 <= item.x1 and item.y0 <= item.y1
    assert x0 <= x1 and y0 <= y1

    # does most of the item area overlap the box?
    # http://math.stackexchange.com/questions/99565/simplest-way-to-calculate-the-intersect-area-of-two-rectangles
    x_overlap = max(0, min(item.x1, x1) - max(item.x0, x0))
    y_overlap = max(0, min(item.y1, y1) - max(item.y0, y0))
    overlap_area = x_overlap * y_overlap
    item_area = (item.x1 - item.x0) * (item.y1 - item.y0)
    assert overlap_area <= item_area

    if DEBUG_BOXHIT and overlap_area != 0:
        print("'%s' %f-%f,%f-%f in %f-%f,%f-%f %2.0f%%" %
              (item.get_text(), item.x0, item.x1, item.y0, item.y1, x0, x1, y0, y1,
               100 * overlap_area / item_area))

    if item_area == 0:
        return False
    else:
        return overlap_area >= 0.5 * item_area

class RectExtractor(TextConverter):
    def __init__(self, rsrcmgr, codec='utf-8', pageno=1, laparams=None):
        dummy = io.StringIO()
        TextConverter.__init__(self, rsrcmgr, outfp=dummy, codec=codec, pageno=pageno, laparams=laparams)
        self.annots = set()

    def setannots(self, annots):
        self.annots = {a for a in annots if a.boxes}

    # main callback from parent PDFConverter
    def receive_layout(self, ltpage):
        self._lasthit = frozenset()
        self._curline = set()
        self.render(ltpage)

    def testboxes(self, item):
        hits = frozenset({a for a in self.annots if any({boxhit(item, b) for b in a.boxes})})
        self._lasthit = hits
        self._curline.update(hits)
        return hits

    # "broadcast" newlines to _all_ annotations that received any text on the
    # current line, in case they see more text on the next line, even if the
    # most recent character was not covered.
    def capture_newline(self):
        for a in self._curline:
            a.capture('\n')
        self._curline = set()

    def render(self, item):
        # If it's a container, recurse on nested items.
        if isinstance(item, LTContainer):
            for child in item:
                self.render(child)

            # Text boxes are a subclass of container, and somehow encode newlines
            # (this weird logic is derived from pdfminer.converter.TextConverter)
            if isinstance(item, LTTextBox):
                self.testboxes(item)
                self.capture_newline()

        # Each character is represented by one LTChar, and we must handle
        # individual characters (not higher-level objects like LTTextLine)
        # so that we can capture only those covered by the annotation boxes.
        elif isinstance(item, LTChar):
            for a in self.testboxes(item):
                a.capture(item.get_text())

        # Annotations capture whitespace not explicitly encoded in
        # the text. They don't have an (X,Y) position, so we need some
        # heuristics to match them to the nearby annotations.
        elif isinstance(item, LTAnno):
            text = item.get_text()
            if text == '\n':
                self.capture_newline()
            else:
                for a in self._lasthit:
                    a.capture(text)


class Page:
    def __init__(self, pageno, mediabox):
        self.pageno = pageno
        self.mediabox = mediabox
        self.annots = []

    def __eq__(self, other):
        return self.pageno == other.pageno

    def __lt__(self, other):
        return self.pageno < other.pageno


class Annotation:
    def __init__(self, page, tagname, coords=None, rect=None, contents=None, author=None):
        self.page = page
        self.tagname = tagname
        if contents == '':
            self.contents = None
        else:
            self.contents = contents
        self.rect = rect
        self.author = author
        self.text = ''

        if coords is None:
            self.boxes = None
        else:
            assert len(coords) % 8 == 0
            self.boxes = []
            while coords != []:
                (x0,y0,x1,y1,x2,y2,x3,y3) = coords[:8]
                coords = coords[8:]
                xvals = [x0, x1, x2, x3]
                yvals = [y0, y1, y2, y3]
                box = (min(xvals), min(yvals), max(xvals), max(yvals))
                self.boxes.append(box)

    def capture(self, text):
        if text == '\n':
            # Kludge for latex: elide hyphens
            if self.text.endswith('-'):
                self.text = self.text[:-1]

            # Join lines, treating newlines as space, while ignoring successive
            # newlines. This makes it easier for the for the renderer to
            # "broadcast" LTAnno newlines to active annotations regardless of
            # box hits. (Detecting paragraph breaks is tricky anyway, and left
            # for future future work!)
            elif not self.text.endswith(' '):
                self.text += ' '
        else:
            self.text += text

    def gettext(self):
        if self.boxes:
            if self.text:
                # replace tex ligatures (and other common odd characters)
                return ''.join([SUBSTITUTIONS.get(c, c) for c in self.text.strip()])
            else:
                # something's strange -- we have boxes but no text for them
                return "(XXX: missing text!)"
        else:
            return None

    def getstartpos(self):
        if self.rect:
            (x0, y0, x1, y1) = self.rect
        elif self.boxes:
            (x0, y0, x1, y1) = self.boxes[0]
        else:
            return None
        # XXX: assume left-to-right top-to-bottom text
        return Pos(self.page, min(x0, x1), max(y0, y1))

    # custom < operator for sorting
    def __lt__(self, other):
        return self.getstartpos() < other.getstartpos()


class Pos:
    def __init__(self, page, x, y):
        self.page = page
        self.x = x
        self.y = y

    def __lt__(self, other):
        if self.page < other.page:
            return True
        elif self.page == other.page:
            assert self.page is other.page
            # XXX: assume left-to-right top-to-bottom documents
            (sx, sy) = self.normalise_to_mediabox()
            (ox, oy) = other.normalise_to_mediabox()
            (x0, y0, x1, y1) = self.page.mediabox
            colwidth = (x1 - x0) / COLUMNS_PER_PAGE
            self_col = (sx - x0) // colwidth
            other_col = (ox - x0) // colwidth
            return self_col < other_col or (self_col == other_col and sy > oy)
        else:
            return False

    def normalise_to_mediabox(self):
        x, y = self.x, self.y
        (x0, y0, x1, y1) = self.page.mediabox
        if x < x0:
            x = x0
        elif x > x1:
            x = x1
        if y < y0:
            y = y0
        elif y > y1:
            y = y1
        return (x, y)


def getannots(pdfannots, page):
    annots = []
    for pa in pdfannots:
        subtype = pa.get('Subtype')
        if subtype is not None and subtype.name not in ANNOT_SUBTYPES:
            continue

        contents = pa.get('Contents')
        if contents is not None:
            # decode as string, normalise line endings, replace special characters
            contents = pdfminer.utils.decode_text(contents)
            contents = contents.replace('\r\n', '\n').replace('\r', '\n')
            contents = ''.join([SUBSTITUTIONS.get(c, c) for c in contents])

        coords = pdftypes.resolve1(pa.get('QuadPoints'))
        rect = pdftypes.resolve1(pa.get('Rect'))
        author = pdftypes.resolve1(pa.get('T'))
        if author is not None:
            author = pdfminer.utils.decode_text(author)
        a = Annotation(page, subtype.name, coords, rect, contents, author=author)
        annots.append(a)

    return annots

def process_file(fh):
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = RectExtractor(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    parser = PDFParser(fh)
    doc = PDFDocument(parser)

    pageslist = [] # pages in page order
    pagesdict = {} # map from PDF page object ID to Page object
    allannots = []

    for (pageno, pdfpage) in enumerate(PDFPage.create_pages(doc)):
        page = Page(pageno, pdfpage.mediabox)
        pageslist.append(page)
        pagesdict[pdfpage.pageid] = page
        if pdfpage.annots:
            pdfannots = []
            for a in pdftypes.resolve1(pdfpage.annots):
                if isinstance(a, pdftypes.PDFObjRef):
                    pdfannots.append(a.resolve())
                else:
                    sys.stderr.write('Warning: unknown annotation: %s\n' % a)

            page.annots = getannots(pdfannots, page)
            page.annots.sort()
            device.setannots(page.annots)
            interpreter.process_page(pdfpage)
            allannots.extend(page.annots)
    device.close()
    return allannots

def create_docx(annotations, filename):
    document = Document()
    table = document.add_table(rows=(len(annotations) + 1), cols=3)
    table.style = 'Table Grid'
    set_col_widths(table)
    heading_cells = table.rows[0].cells
    heading_cells[1].text = 'Editor/Reviewer Comments'
    heading_cells[2].text = 'Response'
    reviewer = ""
    counter = 1
    for key, value in enumerate(annotations):
        if value.tagname == "Underline":
            rawtext = value.gettext()
            text = " ".join(rawtext.strip().split()) if rawtext else ""
            comment = value.contents if value.contents else ""
            row = table.rows[key + 1].cells
            merged = row[1].merge(row[2])
            merged.text = text.capitalize() + " comments"
            shading_elm_0 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            row[0]._tc.get_or_add_tcPr().append(shading_elm_0)
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            merged._tc.get_or_add_tcPr().append(shading_elm_1)
            if comment:
                reviewer = comment
            else:
                reviewer = text
            counter = 1
        else:
            rawtext = value.gettext()
            text = " ".join(rawtext.strip().split()) if rawtext else ""
            comment = value.contents if value.contents else ""
            row = table.rows[key + 1].cells
            if reviewer:
                row[0].text = reviewer + "." + str(counter)
            else:
                row[0].text = str(counter)
            row[1].text = text
            row[2].text = comment
            counter += 1
    document.save(filename)

def create_md(annotations, filename):
    with open(filename, "w", encoding='utf-8') as md_file:
        md_file.write("| | Editor/Reviewer Comments | Response |\n")
        md_file.write("| - | - | - |\n")
        reviewer = ""
        counter = 1
        for key, value in enumerate(annotations):
            if value.tagname == "Underline":
                rawtext = value.gettext()
                text = " ".join(rawtext.strip().split()) if rawtext else ""
                comment = value.contents if value.contents else ""
                md_file.write(f"| | **{text.capitalize()} comments** | |\n")
                if comment:
                    reviewer = comment
                else:
                    reviewer = text
                counter = 1
            else:
                rawtext = value.gettext()
                text = " ".join(rawtext.strip().split()) if rawtext else ""
                comment = value.contents if value.contents else ""
                if reviewer:
                    md_file.write(f"| {reviewer}.{str(counter)} | {text} | {comment} |\n")
                else:
                    md_file.write(f"| {str(counter)} | {text} | {comment} |\n")
                counter += 1

def create_csv(annotations, filename):
    with open(filename, "w", encoding='utf-8', newline='') as csv_file:
        csvwriter = csv.writer(csv_file, delimiter=',', quotechar='"', quoting=csv.QUOTE_ALL)
        csvwriter.writerow(['','Editor/Reviewer Comments', 'Response'])
        reviewer = ""
        counter = 1
        for key, value in enumerate(annotations):
            if value.tagname == "Underline":
                rawtext = value.gettext()
                text = " ".join(rawtext.strip().split()) if rawtext else ""
                comment = value.contents if value.contents else ""
                csvwriter.writerow(['', f'**{text.capitalize()} comments**', ''])
                if comment:
                    reviewer = comment
                else:
                    reviewer = text
                counter = 1
            else:
                rawtext = value.gettext()
                text = " ".join(rawtext.strip().split()) if rawtext else ""
                comment = value.contents if value.contents else ""
                if reviewer:
                    csvwriter.writerow([f'{reviewer}.{str(counter)}', f'{text}', f'{comment}'])
                else:
                    csvwriter.writerow([f'{str(counter)}', f'{text}', f'{comment}'])
                counter += 1

def set_col_widths(table):
    widths = (Inches(0.32), Inches(4.5), Inches(4.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def parse_args():
    p = argparse.ArgumentParser(description=__doc__)

    p.add_argument("input", metavar="INFILE", type=argparse.FileType("rb"),
                   help="PDF files to process", nargs='+')
    p.add_argument("-n", "--cols", default=2, type=int, metavar="COLS", dest="cols",
                   help="number of columns per page in the document (default: 2)")

    return p.parse_args()

def main():
    args = parse_args()

    global COLUMNS_PER_PAGE
    COLUMNS_PER_PAGE = args.cols

    for file in args.input:
        annots = process_file(file)
        create_docx(annots, file.name[:-4] + ".docx")
        create_md(annots, file.name[:-4] + ".md")
        create_csv(annots, file.name[:-4] + ".csv")
    return 0


if __name__ == "__main__":
    sys.exit(main())
